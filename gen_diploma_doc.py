"""
Генерация Word-документа с листингами дипломной работы.
Запуск: python gen_diploma_doc.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_LINE_SPACING

FONT_NAME = 'Consolas'
FONT_SIZE = Pt(8)


def _set_para_spacing(para):
    """Убираем пространство до/после абзаца, одинарный межстрочный интервал."""
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_text(doc, text, bold=False, color=None):
    """Добавить обычный абзац текста шрифтом Consolas 8."""
    para = doc.add_paragraph()
    _set_para_spacing(para)
    run = para.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Устанавливаем восточный шрифт тоже, чтобы кириллица рендерилась корректно
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)
    rPr.insert(0, rFonts)
    return para


def add_heading(doc, text, level=2):
    """Заголовок раздела стилизованный под диплом."""
    para = doc.add_paragraph()
    _set_para_spacing(para)
    para.paragraph_format.space_before = Pt(6)
    run = para.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.font.bold = True
    return para


def add_code_block(doc, code: str):
    """Добавить блок кода построчно."""
    for line in code.splitlines():
        para = doc.add_paragraph()
        _set_para_spacing(para)
        run = para.add_run(line if line else ' ')
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
    # Пустая строка после блока
    add_text(doc, '')


def build_doc():
    doc = Document()

    # Поля страницы
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # -----------------------------------------------------------------------
    # ГЛАВА (фрагмент)
    # -----------------------------------------------------------------------
    add_heading(doc, '3 ТЕХНОЛОГИЧЕСКИЙ РАЗДЕЛ', level=1)
    add_text(doc, '')
    add_heading(doc, '3.1 Описание программной реализации', level=2)
    add_text(doc, '')

    add_text(doc,
        'Приложение Mirmex HelpDesk реализовано на языке программирования Python 3.12 с '
        'использованием веб-фреймворка Django 5.2. Структура проекта соответствует '
        'архитектурному шаблону MVT (Model-View-Template), принятому в Django. '
        'Роль Model выполняет слой данных (models.py), View — слой бизнес-логики (views.py), '
        'Template — слой представления (HTML-шаблоны на основе шаблонизатора Django).')
    add_text(doc, '')
    add_text(doc,
        'Общая структура проекта представлена следующим образом: корневой каталог содержит '
        'пакет mirmex_site (настройки, маршруты, WSGI/ASGI-точки входа) и приложение tickets, '
        'реализующее всю предметную логику системы — модели данных, формы, представления и '
        'HTML-шаблоны.')
    add_text(doc, '')

    # --- settings.py ---
    add_heading(doc, 'Файл mirmex_site/settings.py (конфигурация проекта)', level=3)
    add_text(doc, '')
    add_text(doc,
        'Конфигурационный файл проекта описывает ключевые параметры приложения: '
        'секретный ключ, режим отладки, список разрешённых хостов, подключённые '
        'приложения, стек промежуточных обработчиков (middleware), параметры базы '
        'данных и раздачи статических файлов. '
        'Чувствительные настройки (SECRET_KEY, параметры БД) вынесены в переменные '
        'окружения и считываются через библиотеку python-decouple, что обеспечивает '
        'безопасность при развёртывании. '
        'База данных PostgreSQL подключается либо через переменную DATABASE_URL '
        '(для облачного хостинга Railway), либо через отдельные переменные DB_NAME, '
        'DB_USER, DB_PASSWORD, DB_HOST, DB_PORT. '
        'Раздача статических файлов в production обеспечивается промежуточным '
        'обработчиком WhiteNoise.')
    add_text(doc, '')
    add_text(doc, 'Ниже приведён листинг файла настроек:')
    add_text(doc, '')

    settings_code = '''\
"""
Django settings for mirmex_site project.
"""

from pathlib import Path
from decouple import config, Csv
import dj_database_url

BASE_DIR = Path(__file__).resolve().parent.parent

SECRET_KEY = config('SECRET_KEY')
DEBUG = config('DEBUG', default=False, cast=bool)
ALLOWED_HOSTS = config('ALLOWED_HOSTS', default='localhost,127.0.0.1', cast=Csv())

INSTALLED_APPS = [
    "tickets",
    "django.contrib.admin",
    "django.contrib.auth",
    "django.contrib.contenttypes",
    "django.contrib.sessions",
    "django.contrib.messages",
    "django.contrib.staticfiles",
]

MIDDLEWARE = [
    "django.middleware.security.SecurityMiddleware",
    "whitenoise.middleware.WhiteNoiseMiddleware",
    "django.contrib.sessions.middleware.SessionMiddleware",
    "django.middleware.common.CommonMiddleware",
    "django.middleware.csrf.CsrfViewMiddleware",
    "django.contrib.auth.middleware.AuthenticationMiddleware",
    "django.contrib.messages.middleware.MessageMiddleware",
    "django.middleware.clickjacking.XFrameOptionsMiddleware",
]

ROOT_URLCONF = "mirmex_site.urls"

TEMPLATES = [{
    "BACKEND": "django.template.backends.django.DjangoTemplates",
    "DIRS": [],
    "APP_DIRS": True,
    "OPTIONS": {
        "context_processors": [
            "django.template.context_processors.request",
            "django.contrib.auth.context_processors.auth",
            "django.contrib.messages.context_processors.messages",
            "tickets.context_processors.user_role",
        ],
    },
}]

WSGI_APPLICATION = "mirmex_site.wsgi.application"

LOGIN_URL = '/accounts/login/'
LOGOUT_REDIRECT_URL = '/accounts/login/'
LOGIN_REDIRECT_URL = '/tickets/'

_db_url = config('DATABASE_URL', default=None)
if _db_url:
    DATABASES = {'default': dj_database_url.parse(_db_url, conn_max_age=600)}
else:
    DATABASES = {
        'default': {
            'ENGINE': 'django.db.backends.postgresql',
            'NAME': config('DB_NAME', default='mirmex_helpdesk'),
            'USER': config('DB_USER', default='postgres'),
            'PASSWORD': config('DB_PASSWORD', default='postgres'),
            'HOST': config('DB_HOST', default='127.0.0.1'),
            'PORT': config('DB_PORT', default='5432'),
        }
    }

LANGUAGE_CODE = "ru-ru"
TIME_ZONE = "Europe/Moscow"
USE_I18N = True
USE_TZ = True

STATIC_URL = '/static/'
STATIC_ROOT = BASE_DIR / 'staticfiles'
STATICFILES_DIRS = [BASE_DIR / 'mirmex_site' / 'static']
STORAGES = {
    "default": {"BACKEND": "django.core.files.storage.FileSystemStorage"},
    "staticfiles": {
        "BACKEND": "whitenoise.storage.CompressedManifestStaticFilesStorage"
    },
}

DEFAULT_AUTO_FIELD = "django.db.models.BigAutoField"
'''
    add_code_block(doc, settings_code)

    # --- urls.py ---
    add_heading(doc, 'Файл mirmex_site/urls.py (маршрутизация)', level=3)
    add_text(doc, '')
    add_text(doc,
        'Файл маршрутизации верхнего уровня связывает URL-адреса с соответствующими '
        'обработчиками. Основные маршруты приложения инкапсулированы в модуле '
        'tickets.urls и подключаются через механизм include. Дополнительно описаны '
        'маршруты встроенной системы аутентификации Django (вход, выход, смена пароля) '
        'и маршрут самостоятельной регистрации пользователей.')
    add_text(doc, '')

    urls_main_code = '''\
from django.contrib import admin
from django.urls import path, include
from django.contrib.auth import views as auth_views
from django.views.generic import RedirectView
from tickets.views import RegisterView

urlpatterns = [
    path("admin/", admin.site.urls),
    path("", RedirectView.as_view(url='/tickets/', permanent=False)),
    path("", include('tickets.urls')),

    path('accounts/login/',  auth_views.LoginView.as_view(),  name='login'),
    path('accounts/logout/', auth_views.LogoutView.as_view(), name='logout'),
    path('accounts/register/', RegisterView.as_view(),        name='register'),

    path('accounts/password_change/',
         auth_views.PasswordChangeView.as_view(
             template_name='registration/password_change.html',
             success_url='/profile/'),
         name='password_change'),
    path('accounts/password_change/done/',
         auth_views.PasswordChangeDoneView.as_view(
             template_name='registration/password_change_done.html'),
         name='password_change_done'),
]

handler403 = 'django.views.defaults.permission_denied'
handler404 = 'django.views.defaults.page_not_found'
'''
    add_code_block(doc, urls_main_code)

    # --- models.py ---
    add_heading(doc, 'Файл tickets/models.py (модели данных)', level=3)
    add_text(doc, '')
    add_text(doc,
        'Модели данных являются центральным элементом приложения и описывают структуру '
        'реляционной базы данных. В системе определены четыре модели.')
    add_text(doc, '')
    add_text(doc,
        'Модель Equipment описывает IT-оборудование организации. Каждая единица '
        'оборудования имеет серийный номер (уникальный), модель, тип (рабочая станция, '
        'сервер, сеть, периферия, ПО, мобильное устройство), местоположение, статус '
        '(в использовании, в ремонте, на складе), даты приобретения и окончания гарантии. '
        'Методы warranty_expired() и warranty_expiring_soon() реализуют логику '
        'определения состояния гарантийного срока.')
    add_text(doc, '')
    add_text(doc,
        'Модель Tickets является главной сущностью системы. Заявка связана с '
        'оборудованием (необязательно), заявителем и техником через внешние ключи. '
        'Жизненный цикл заявки описан допустимыми переходами между статусами: '
        'новая → назначена → в работе → закрыта. '
        'Переходы контролируются словарём ALLOWED_TRANSITIONS и методами assign(), '
        'start_work(), change_status(), close(). Метод is_overdue() позволяет '
        'определить, нарушен ли SLA-срок. Метод resolution_hours() вычисляет '
        'фактическое время решения в часах.')
    add_text(doc, '')
    add_text(doc,
        'Модель Comment хранит комментарии пользователей к заявкам. '
        'Модель TicketHistory реализует полный аудит: каждое действие над заявкой '
        '(создание, смена статуса, назначение, комментарий, редактирование) фиксируется '
        'отдельной записью с указанием пользователя и временной метки.')
    add_text(doc, '')
    add_text(doc, 'Ниже приведён листинг ключевых частей файла models.py:')
    add_text(doc, '')

    models_excerpt = '''\
class Equipment(models.Model):
    STATUS_CHOICES = [
        ('in_use',    'В использовании'),
        ('in_repair', 'В ремонте'),
        ('storage',   'На складе'),
    ]
    TYPE_CHOICES = [
        ('workstation', 'Рабочая станция'),
        ('server',      'Сервер'),
        ('network',     'Сеть'),
        ('peripheral',  'Периферийное устройство'),
        ('software',    'Программное обеспечение'),
        ('mobile',      'Мобильное устройство'),
        ('other',       'Другое'),
    ]

    serial         = models.CharField(max_length=100, unique=True)
    model          = models.CharField(max_length=200)
    equipment_type = models.CharField(max_length=20, choices=TYPE_CHOICES, default='other')
    location       = models.CharField(max_length=200, blank=True)
    status         = models.CharField(max_length=20, choices=STATUS_CHOICES, default='in_use')
    purchased_at   = models.DateField(null=True, blank=True)
    warranty_until = models.DateField(null=True, blank=True)
    notes          = models.TextField(blank=True)

    def warranty_expired(self):
        if self.warranty_until:
            return timezone.now().date() > self.warranty_until
        return False

    def warranty_expiring_soon(self):
        if self.warranty_until and not self.warranty_expired():
            return (self.warranty_until - timezone.now().date()).days <= 30
        return False


class Tickets(models.Model):
    STATUS_CHOICES = [
        ('new',         'Новая'),
        ('assigned',    'Назначена'),
        ('in_progress', 'В работе'),
        ('closed',      'Закрыта'),
    ]
    PRIORITY_CHOICES = [
        ('low',      'Низкий'),
        ('medium',   'Средний'),
        ('high',     'Высокий'),
        ('critical', 'Критический'),
    ]
    CATEGORY_CHOICES = [
        ('incident',        'Инцидент'),
        ('service_request', 'Запрос на обслуживание'),
        ('consultation',    'Консультация'),
        ('change_request',  'Запрос на изменение'),
        ('admin_request',   'Запрос к администратору'),
        ('other',           'Другое'),
    ]
    ALLOWED_TRANSITIONS = {
        'new':         ['assigned', 'closed'],
        'assigned':    ['in_progress', 'closed'],
        'in_progress': ['closed'],
        'closed':      [],
    }

    equipment  = models.ForeignKey(Equipment, on_delete=models.SET_NULL,
                                   null=True, blank=True, related_name='tickets')
    reporter   = models.ForeignKey(User, related_name='reported_tickets',
                                   on_delete=models.SET_NULL, null=True)
    technician = models.ForeignKey(User, related_name='assigned_tickets',
                                   on_delete=models.SET_NULL, null=True, blank=True)
    title      = models.CharField(max_length=255)
    description= models.TextField()
    priority   = models.CharField(max_length=20, choices=PRIORITY_CHOICES, default='medium')
    category   = models.CharField(max_length=20, choices=CATEGORY_CHOICES, default='other')
    due_date   = models.DateTimeField(null=True, blank=True)
    created_at = models.DateTimeField(auto_now_add=True)
    status     = models.CharField(max_length=20, choices=STATUS_CHOICES, default='new')
    closed_at  = models.DateTimeField(null=True, blank=True)

    def assign(self, technician, changed_by=None, bypass_validation=False):
        if not technician:
            raise ValidationError("Не указан техник")
        self.technician = technician
        self.change_status('assigned', changed_by=changed_by,
                           bypass_validation=bypass_validation)
        self.save(update_fields=['technician', 'status', 'closed_at'])
        TicketHistory.objects.create(
            ticket=self, changed_by=changed_by,
            action='assigned',
            comment=f'Назначен техник: '
                    f'{technician.get_full_name() or technician.username}',
        )

    def start_work(self, changed_by=None, bypass_validation=False):
        if not self.technician:
            raise ValidationError("Нельзя начать работу без техника")
        self.change_status('in_progress', changed_by=changed_by,
                           bypass_validation=bypass_validation)

    def change_status(self, new_status, changed_by=None, bypass_validation=False):
        old_status = self.status
        if not bypass_validation:
            allowed = self.ALLOWED_TRANSITIONS.get(self.status, [])
            if new_status not in allowed:
                raise ValidationError(
                    f"Недопустимый переход: {self.status} -> {new_status}")
        self.status = new_status
        if new_status == 'closed' and self.closed_at is None:
            self.closed_at = timezone.now()
        self.save()
        if old_status != new_status:
            status_display = dict(self.STATUS_CHOICES)
            TicketHistory.objects.create(
                ticket=self, changed_by=changed_by,
                action='status_changed',
                old_status=old_status, new_status=new_status,
                comment=f'{status_display.get(old_status)} -> '
                        f'{status_display.get(new_status)}',
            )

    def close(self, changed_by=None, bypass_validation=False):
        self.change_status('closed', changed_by=changed_by,
                           bypass_validation=bypass_validation)

    def is_overdue(self):
        if self.due_date and self.status != 'closed':
            return timezone.now() > self.due_date
        return False

    def resolution_hours(self):
        if self.closed_at:
            delta = self.closed_at - self.created_at
            return round(delta.total_seconds() / 3600, 1)
        return None
'''
    add_code_block(doc, models_excerpt)

    # --- forms.py ---
    add_heading(doc, 'Файл tickets/forms.py (формы ввода данных)', level=3)
    add_text(doc, '')
    add_text(doc,
        'Формы Django используются для валидации и обработки пользовательского ввода. '
        'В системе реализованы следующие формы: TicketCreateForm — для создания новой '
        'заявки (поля: оборудование, тема, описание, приоритет, категория, срок SLA); '
        'TicketUpdateForm — для изменения заявки с разграничением прав (суперпользователь '
        'редактирует все поля, техник — только статус, диспетчер — все поля, кроме '
        'прямых правок модели); CommentForm — текстовый комментарий к заявке; '
        'EquipmentForm — карточка оборудования; UserEditForm — управление ролью и '
        'данными пользователя администратором; UserRegistrationForm — форма '
        'самостоятельной регистрации с двойным вводом пароля и валидацией по политике '
        'безопасности Django; ProfileForm — редактирование собственного профиля.')
    add_text(doc, '')
    add_text(doc, 'Пример: форма создания заявки TicketCreateForm:')
    add_text(doc, '')

    forms_excerpt = '''\
class TicketCreateForm(forms.ModelForm):
    class Meta:
        model = Tickets
        fields = ['equipment', 'title', 'description', 'priority', 'category', 'due_date']
        widgets = {
            'due_date': forms.DateTimeInput(
                attrs={'type': 'datetime-local'}, format='%Y-%m-%dT%H:%M'),
        }

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.fields['due_date'].input_formats = ['%Y-%m-%dT%H:%M']
        self.fields['equipment'].queryset = Equipment.objects.all()
        self.fields['equipment'].empty_label = \
            '--- Выберите оборудование (необязательно) ---'
'''
    add_code_block(doc, forms_excerpt)

    add_text(doc, 'Пример: форма обновления заявки TicketUpdateForm (фрагмент):')
    add_text(doc, '')

    ticket_update_excerpt = '''\
class TicketUpdateForm(forms.ModelForm):
    class Meta:
        model = Tickets
        fields = ['status', 'technician', 'priority', 'category', 'due_date']

    def __init__(self, *args, **kwargs):
        self.user = kwargs.pop('user', None)
        super().__init__(*args, **kwargs)
        self.original_status = self.instance.status
        # По умолчанию все поля заблокированы
        for f in self.fields.values():
            f.disabled = True
        if self.user and self.user.is_superuser:
            for f in self.fields.values():
                f.disabled = False
        elif self.user and self.user.groups.filter(name='technician').exists():
            self.fields['status'].disabled = False
        self.fields['technician'].queryset = \
            User.objects.filter(groups__name='technician')

    def clean(self):
        cleaned_data = super().clean()
        new_status = cleaned_data.get('status')
        if self.user and not self.user.is_superuser:
            if new_status and new_status != self.original_status:
                allowed = self.instance.ALLOWED_TRANSITIONS.get(
                    self.original_status, [])
                if new_status not in allowed:
                    raise ValidationError({'status':
                        f"Недопустимый переход из "
                        f"«{self.instance.get_status_display()}»."})
        return cleaned_data
'''
    add_code_block(doc, ticket_update_excerpt)

    # --- views.py ---
    add_heading(doc, 'Файл tickets/views.py (представления / бизнес-логика)', level=3)
    add_text(doc, '')
    add_text(doc,
        'Файл views.py содержит всю бизнес-логику приложения. В начале файла объявлены '
        'вспомогательные функции определения роли пользователя: is_admin_or_dispatcher(), '
        'is_technician(), get_user_role(). Они используются во всех представлениях для '
        'ограничения доступа и формирования контекста.')
    add_text(doc, '')
    add_text(doc,
        'TicketListView реализует отображение списка заявок с фильтрацией по статусу, '
        'приоритету, категории, технику и полнотекстовым поиском. Queryset формируется '
        'динамически в зависимости от роли: администратор и диспетчер видят все заявки, '
        'техник — только назначенные ему, заявитель — только свои. Реализована '
        'постраничная навигация (15 записей на страницу).')
    add_text(doc, '')
    add_text(doc, 'Фрагмент TicketListView:')
    add_text(doc, '')

    ticket_list_excerpt = '''\
class TicketListView(LoginRequiredMixin, ListView):
    model = Tickets
    template_name = 'tickets/ticket_list.html'
    paginate_by = 15

    def get_queryset(self):
        user = self.request.user
        if is_admin_or_dispatcher(user):
            queryset = Tickets.objects.select_related(
                'reporter', 'technician', 'equipment')
        elif is_technician(user):
            queryset = Tickets.objects.filter(
                technician=user).select_related(
                'reporter', 'technician', 'equipment')
        else:
            queryset = Tickets.objects.filter(
                reporter=user).select_related(
                'reporter', 'technician', 'equipment')

        status_filter = self.request.GET.get('status')
        if status_filter:
            queryset = queryset.filter(status=status_filter)

        priority_filter = self.request.GET.get('priority')
        if priority_filter:
            queryset = queryset.filter(priority=priority_filter)

        category_filter = self.request.GET.get('category')
        if category_filter:
            queryset = queryset.filter(category=category_filter)

        search_query = self.request.GET.get('q', '').strip()
        if search_query:
            queryset = queryset.filter(
                Q(title__icontains=search_query) |
                Q(description__icontains=search_query) |
                Q(id__icontains=search_query))

        return queryset.order_by('-created_at')
'''
    add_code_block(doc, ticket_list_excerpt)

    add_text(doc,
        'TicketDetailView отображает карточку заявки с историей изменений, '
        'комментариями и индикатором прогресса SLA. Вычисление SLA выполняется '
        'путём сравнения текущего времени с датой создания заявки и сроком due_date. '
        'Результат передаётся в шаблон в виде числового процента и цвета индикатора.')
    add_text(doc, '')
    add_text(doc, 'Фрагмент вычисления SLA из TicketDetailView:')
    add_text(doc, '')

    sla_excerpt = '''\
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        ticket = self.object

        sla_percent = None
        sla_color = 'success'
        if ticket.due_date and ticket.status != ticket.STATUS_CLOSED:
            now = timezone.now()
            total_seconds   = (ticket.due_date - ticket.created_at).total_seconds()
            elapsed_seconds = (now - ticket.created_at).total_seconds()
            if total_seconds > 0:
                sla_percent = min(100,
                    round(elapsed_seconds / total_seconds * 100))
                if   sla_percent >= 100: sla_color = 'danger'
                elif sla_percent >= 75:  sla_color = 'warning'
                elif sla_percent >= 50:  sla_color = 'info'
                else:                    sla_color = 'success'
        context['sla_percent'] = sla_percent
        context['sla_color']   = sla_color
        ...
        return context
'''
    add_code_block(doc, sla_excerpt)

    add_text(doc,
        'ReportsView (доступна только администратору и диспетчеру) собирает '
        'агрегированную статистику: общее количество заявок, количество открытых '
        'и закрытых, созданных и закрытых за текущий месяц, просроченных. '
        'Данные по категориям, статусам и приоритетам подготавливаются для '
        'Chart.js и передаются в шаблон в формате JSON. Также формируется '
        'таблица эффективности технических специалистов с показателями '
        'общего числа заявок, количества открытых и среднего времени решения.')
    add_text(doc, '')
    add_text(doc, 'Функции управления действиями над заявкой (assign_ticket, start_ticket, '
        'close_ticket, add_comment) реализованы как функциональные представления '
        'с декоратором @login_required. Каждая из них проверяет права доступа и '
        'при успешном выполнении действия создаёт запись в TicketHistory.')
    add_text(doc, '')

    # --- urls.py (app) ---
    add_heading(doc, 'Файл tickets/urls.py (маршруты приложения)', level=3)
    add_text(doc, '')
    add_text(doc,
        'Маршруты приложения tickets описывают связь URL-адресов с соответствующими '
        'классами-представлениями и функциями. Структура URL-пространства разделена '
        'на логические группы: оборудование, заявки, Kanban-доска, поиск, отчёты, '
        'профиль, управление пользователями, страница "О программном продукте".')
    add_text(doc, '')

    tickets_urls_code = '''\
from django.urls import path
from . import views

urlpatterns = [
    # Оборудование
    path('equipment/',              views.EquipmentListView.as_view(),   name='equipment_list'),
    path('equipment/create/',       views.EquipmentCreateView.as_view(), name='equipment_create'),
    path('equipment/<int:pk>/',     views.EquipmentDetailView.as_view(), name='equipment_detail'),
    path('equipment/<int:pk>/edit/',views.EquipmentUpdateView.as_view(), name='equipment_edit'),

    # Заявки
    path('tickets/',                   views.TicketListView.as_view(),      name='ticket_list'),
    path('tickets/export/',            views.TicketExportCSVView.as_view(), name='ticket_export'),
    path('tickets/create/',            views.TicketCreateView.as_view(),    name='ticket_create'),
    path('tickets/<int:pk>/',          views.TicketDetailView.as_view(),    name='ticket_detail'),
    path('tickets/<int:pk>/edit/',     views.TicketUpdateView.as_view(),    name='ticket_update'),
    path('tickets/<int:pk>/assign/',   views.assign_ticket,                 name='ticket_assign'),
    path('tickets/<int:pk>/start/',    views.start_ticket,                  name='ticket_start'),
    path('tickets/<int:pk>/close/',    views.close_ticket,                  name='ticket_close'),
    path('tickets/<int:pk>/comment/',  views.add_comment,                   name='ticket_comment'),

    # Kanban-доска
    path('kanban/',   views.KanbanBoardView.as_view(), name='kanban'),

    # Поиск
    path('search/',   views.SearchView.as_view(),      name='search'),

    # Отчёты
    path('reports/',  views.ReportsView.as_view(),     name='reports'),

    # Профиль
    path('profile/',  views.UserProfileView.as_view(), name='profile'),

    # Управление пользователями
    path('users/',              views.UserManagementView.as_view(), name='user_list'),
    path('users/<int:pk>/edit/',views.user_edit,                    name='user_edit'),

    # О продукте
    path('about/',    views.AboutView.as_view(),       name='about'),
]
'''
    add_code_block(doc, tickets_urls_code)

    # -----------------------------------------------------------------------
    # PAGE BREAK → APPENDIX
    # -----------------------------------------------------------------------
    doc.add_page_break()

    add_heading(doc, 'ПРИЛОЖЕНИЕ', level=1)
    add_text(doc, '')

    # -----------------------------------------------------------------------
    # ПРИЛОЖЕНИЕ — полные листинги
    # -----------------------------------------------------------------------

    # --- полный models.py ---
    add_heading(doc, 'Файл tickets/models.py (полный листинг)', level=2)
    add_text(doc, '')

    models_full = '''\
from django.db import models
from django.contrib.auth import get_user_model
from django.utils import timezone
from django.core.exceptions import ValidationError

User = get_user_model()


class Equipment(models.Model):
    STATUS_IN_USE    = 'in_use'
    STATUS_IN_REPAIR = 'in_repair'
    STATUS_STORAGE   = 'storage'

    STATUS_CHOICES = [
        (STATUS_IN_USE,    'В использовании'),
        (STATUS_IN_REPAIR, 'В ремонте'),
        (STATUS_STORAGE,   'На складе'),
    ]

    TYPE_WORKSTATION = 'workstation'
    TYPE_SERVER      = 'server'
    TYPE_NETWORK     = 'network'
    TYPE_PERIPHERAL  = 'peripheral'
    TYPE_SOFTWARE    = 'software'
    TYPE_MOBILE      = 'mobile'
    TYPE_OTHER       = 'other'

    TYPE_CHOICES = [
        (TYPE_WORKSTATION, 'Рабочая станция'),
        (TYPE_SERVER,      'Сервер'),
        (TYPE_NETWORK,     'Сеть'),
        (TYPE_PERIPHERAL,  'Периферийное устройство'),
        (TYPE_SOFTWARE,    'Программное обеспечение'),
        (TYPE_MOBILE,      'Мобильное устройство'),
        (TYPE_OTHER,       'Другое'),
    ]

    serial         = models.CharField(max_length=100, unique=True,
                                      verbose_name='Серийный номер')
    model          = models.CharField(max_length=200, verbose_name='Модель')
    equipment_type = models.CharField(max_length=20, choices=TYPE_CHOICES,
                                      default=TYPE_OTHER,
                                      verbose_name='Тип оборудования')
    location       = models.CharField(max_length=200, blank=True,
                                      verbose_name='Расположение')
    status         = models.CharField(max_length=20, choices=STATUS_CHOICES,
                                      default=STATUS_IN_USE, verbose_name='Статус')
    purchased_at   = models.DateField(null=True, blank=True,
                                      verbose_name='Дата приобретения')
    warranty_until = models.DateField(null=True, blank=True,
                                      verbose_name='Гарантия до')
    notes          = models.TextField(blank=True, verbose_name='Примечания')

    class Meta:
        verbose_name        = 'Оборудование'
        verbose_name_plural = 'Оборудование'
        ordering = ['model']

    def warranty_expired(self):
        if self.warranty_until:
            return timezone.now().date() > self.warranty_until
        return False

    def warranty_expiring_soon(self):
        if self.warranty_until and not self.warranty_expired():
            days_left = (self.warranty_until - timezone.now().date()).days
            return days_left <= 30
        return False

    def __str__(self):
        return f"{self.model} ({self.serial})"


class Tickets(models.Model):
    STATUS_NEW         = 'new'
    STATUS_ASSIGNED    = 'assigned'
    STATUS_IN_PROGRESS = 'in_progress'
    STATUS_CLOSED      = 'closed'

    STATUS_CHOICES = [
        (STATUS_NEW,         'Новая'),
        (STATUS_ASSIGNED,    'Назначена'),
        (STATUS_IN_PROGRESS, 'В работе'),
        (STATUS_CLOSED,      'Закрыта'),
    ]

    PRIORITY_LOW      = 'low'
    PRIORITY_MEDIUM   = 'medium'
    PRIORITY_HIGH     = 'high'
    PRIORITY_CRITICAL = 'critical'

    PRIORITY_CHOICES = [
        (PRIORITY_LOW,      'Низкий'),
        (PRIORITY_MEDIUM,   'Средний'),
        (PRIORITY_HIGH,     'Высокий'),
        (PRIORITY_CRITICAL, 'Критический'),
    ]

    CATEGORY_INCIDENT        = 'incident'
    CATEGORY_SERVICE_REQUEST = 'service_request'
    CATEGORY_CONSULTATION    = 'consultation'
    CATEGORY_CHANGE_REQUEST  = 'change_request'
    CATEGORY_ADMIN_REQUEST   = 'admin_request'
    CATEGORY_OTHER           = 'other'

    CATEGORY_CHOICES = [
        (CATEGORY_INCIDENT,        'Инцидент'),
        (CATEGORY_SERVICE_REQUEST, 'Запрос на обслуживание'),
        (CATEGORY_CONSULTATION,    'Консультация'),
        (CATEGORY_CHANGE_REQUEST,  'Запрос на изменение'),
        (CATEGORY_ADMIN_REQUEST,   'Запрос к администратору'),
        (CATEGORY_OTHER,           'Другое'),
    ]

    ALLOWED_TRANSITIONS = {
        STATUS_NEW:         [STATUS_ASSIGNED, STATUS_CLOSED],
        STATUS_ASSIGNED:    [STATUS_IN_PROGRESS, STATUS_CLOSED],
        STATUS_IN_PROGRESS: [STATUS_CLOSED],
        STATUS_CLOSED:      [],
    }

    equipment  = models.ForeignKey(Equipment, on_delete=models.SET_NULL,
                                   null=True, blank=True, related_name='tickets',
                                   verbose_name='Оборудование')
    reporter   = models.ForeignKey(User, related_name='reported_tickets',
                                   on_delete=models.SET_NULL, null=True,
                                   verbose_name='Заявитель')
    technician = models.ForeignKey(User, related_name='assigned_tickets',
                                   on_delete=models.SET_NULL, null=True, blank=True,
                                   verbose_name='Техник')
    title       = models.CharField(max_length=255, verbose_name='Тема')
    description = models.TextField(verbose_name='Описание')
    priority    = models.CharField(max_length=20, choices=PRIORITY_CHOICES,
                                   default=PRIORITY_MEDIUM, verbose_name='Приоритет')
    category    = models.CharField(max_length=20, choices=CATEGORY_CHOICES,
                                   default=CATEGORY_OTHER, verbose_name='Категория')
    due_date    = models.DateTimeField(null=True, blank=True,
                                       verbose_name='Срок выполнения (SLA)')
    created_at  = models.DateTimeField(auto_now_add=True, verbose_name='Создана')
    status      = models.CharField(max_length=20, choices=STATUS_CHOICES,
                                   default=STATUS_NEW, verbose_name='Статус')
    closed_at   = models.DateTimeField(null=True, blank=True, verbose_name='Закрыта')

    class Meta:
        verbose_name        = 'Заявка'
        verbose_name_plural = 'Заявки'
        ordering = ['-created_at']
        permissions = [
            ('can_assign_ticket',  'Can assign ticket'),
            ('can_create_ticket',  'Can create ticket'),
            ('can_change_status',  'Can change ticket status'),
        ]

    def assign(self, technician, changed_by=None, bypass_validation=False):
        if not technician:
            raise ValidationError("Не указан техник")
        self.technician = technician
        self.change_status(self.STATUS_ASSIGNED, changed_by=changed_by,
                           bypass_validation=bypass_validation)
        self.save(update_fields=['technician', 'status', 'closed_at'])
        TicketHistory.objects.create(
            ticket=self, changed_by=changed_by,
            action=TicketHistory.ACTION_ASSIGNED,
            comment=f'Назначен техник: '
                    f'{technician.get_full_name() or technician.username}',
        )

    def start_work(self, changed_by=None, bypass_validation=False):
        if not self.technician:
            raise ValidationError("Нельзя начать работу без техника")
        self.change_status(self.STATUS_IN_PROGRESS, changed_by=changed_by,
                           bypass_validation=bypass_validation)

    def change_status(self, new_status, changed_by=None, bypass_validation=False):
        old_status = self.status
        if not bypass_validation:
            allowed = self.ALLOWED_TRANSITIONS.get(self.status, [])
            if new_status not in allowed:
                raise ValidationError(
                    f"Недопустимый переход: {self.status} -> {new_status}")
        self.status = new_status
        if new_status == self.STATUS_CLOSED and self.closed_at is None:
            self.closed_at = timezone.now()
        self.save()
        if old_status != new_status:
            status_display = dict(self.STATUS_CHOICES)
            TicketHistory.objects.create(
                ticket=self, changed_by=changed_by,
                action=TicketHistory.ACTION_STATUS_CHANGED,
                old_status=old_status, new_status=new_status,
                comment=f'{status_display.get(old_status)} -> '
                        f'{status_display.get(new_status)}',
            )

    def close(self, changed_by=None, bypass_validation=False):
        self.change_status(self.STATUS_CLOSED, changed_by=changed_by,
                           bypass_validation=bypass_validation)

    def is_overdue(self):
        if self.due_date and self.status != self.STATUS_CLOSED:
            return timezone.now() > self.due_date
        return False

    def resolution_hours(self):
        if self.closed_at:
            delta = self.closed_at - self.created_at
            return round(delta.total_seconds() / 3600, 1)
        return None

    def __str__(self):
        return f"{self.title} #{self.pk}"


class Comment(models.Model):
    ticket     = models.ForeignKey(Tickets, on_delete=models.CASCADE,
                                   related_name='comments', verbose_name='Заявка')
    author     = models.ForeignKey(User, on_delete=models.SET_NULL,
                                   null=True, verbose_name='Автор')
    text       = models.TextField(verbose_name='Комментарий')
    created_at = models.DateTimeField(auto_now_add=True)

    class Meta:
        verbose_name        = 'Комментарий'
        verbose_name_plural = 'Комментарии'
        ordering = ['created_at']

    def __str__(self):
        return f"Комментарий от {self.author} к #{self.ticket.pk}"


class TicketHistory(models.Model):
    ACTION_CREATED        = 'created'
    ACTION_STATUS_CHANGED = 'status_changed'
    ACTION_ASSIGNED       = 'assigned'
    ACTION_COMMENTED      = 'commented'
    ACTION_EDITED         = 'edited'

    ACTION_CHOICES = [
        (ACTION_CREATED,        'Создана'),
        (ACTION_STATUS_CHANGED, 'Статус изменён'),
        (ACTION_ASSIGNED,       'Назначен техник'),
        (ACTION_COMMENTED,      'Добавлен комментарий'),
        (ACTION_EDITED,         'Отредактирована'),
    ]

    ticket     = models.ForeignKey(Tickets, on_delete=models.CASCADE,
                                   related_name='history', verbose_name='Заявка')
    changed_by = models.ForeignKey(User, on_delete=models.SET_NULL,
                                   null=True, blank=True, verbose_name='Изменил')
    action     = models.CharField(max_length=30, choices=ACTION_CHOICES,
                                  default=ACTION_STATUS_CHANGED, verbose_name='Действие')
    old_status = models.CharField(max_length=20, blank=True)
    new_status = models.CharField(max_length=20, blank=True)
    comment    = models.TextField(blank=True, verbose_name='Описание')
    timestamp  = models.DateTimeField(auto_now_add=True)

    class Meta:
        verbose_name        = 'История заявки'
        verbose_name_plural = 'История заявок'
        ordering = ['timestamp']

    def __str__(self):
        return f"История #{self.ticket.pk} — {self.get_action_display()}"
'''
    add_code_block(doc, models_full)

    # --- полный forms.py ---
    add_heading(doc, 'Файл tickets/forms.py (полный листинг)', level=2)
    add_text(doc, '')

    forms_full = '''\
from django import forms
from django.contrib.auth import get_user_model
from django.contrib.auth.password_validation import validate_password
from django.core.exceptions import ValidationError
from .models import Tickets, Equipment, Comment

User = get_user_model()


class TicketCreateForm(forms.ModelForm):
    class Meta:
        model = Tickets
        fields = ['equipment', 'title', 'description', 'priority', 'category', 'due_date']
        widgets = {
            'due_date': forms.DateTimeInput(
                attrs={'type': 'datetime-local'}, format='%Y-%m-%dT%H:%M'),
        }

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.fields['due_date'].input_formats = ['%Y-%m-%dT%H:%M']
        self.fields['equipment'].queryset = Equipment.objects.all()
        self.fields['equipment'].empty_label = \
            '--- Выберите оборудование (необязательно) ---'


class TicketUpdateForm(forms.ModelForm):
    class Meta:
        model = Tickets
        fields = ['status', 'technician', 'priority', 'category', 'due_date']
        widgets = {
            'due_date': forms.DateTimeInput(
                attrs={'type': 'datetime-local'}, format='%Y-%m-%dT%H:%M'),
        }

    def __init__(self, *args, **kwargs):
        self.user = kwargs.pop('user', None)
        super().__init__(*args, **kwargs)
        self.original_status     = self.instance.status
        self.original_technician = self.instance.technician
        self.fields['due_date'].input_formats = ['%Y-%m-%dT%H:%M']
        for f in self.fields.values():
            f.disabled = True
        if self.user is None:
            return
        if self.user.is_superuser:
            for f in self.fields.values():
                f.disabled = False
        else:
            is_dispatcher = self.user.groups.filter(name='dispatcher').exists()
            is_technician = self.user.groups.filter(name='technician').exists()
            if is_dispatcher:
                for f in self.fields.values():
                    f.disabled = False
            elif is_technician:
                self.fields['status'].disabled = False
        self.fields['technician'].queryset = \
            User.objects.filter(groups__name='technician')

    def clean(self):
        cleaned_data = super().clean()
        new_status = cleaned_data.get('status')
        if self.user and not self.user.is_superuser:
            if new_status and new_status != self.original_status:
                allowed = self.instance.ALLOWED_TRANSITIONS.get(
                    self.original_status, [])
                if new_status not in allowed:
                    raise ValidationError({'status':
                        f"Недопустимый переход из "
                        f"«{self.instance.get_status_display()}» в "
                        f"«{dict(self.instance.STATUS_CHOICES)[new_status]}»."})
        return cleaned_data

    def save(self, commit=True):
        new_status      = self.cleaned_data.get('status')
        new_technician  = self.cleaned_data.get('technician')
        if self.user and self.user.is_superuser:
            self.instance.status     = new_status
            self.instance.technician = new_technician
            self.instance.priority   = self.cleaned_data.get(
                'priority', self.instance.priority)
            self.instance.category   = self.cleaned_data.get(
                'category', self.instance.category)
            self.instance.due_date   = self.cleaned_data.get(
                'due_date', self.instance.due_date)
            if (new_status == self.instance.STATUS_CLOSED
                    and self.instance.closed_at is None):
                from django.utils import timezone
                self.instance.closed_at = timezone.now()
            if commit:
                self.instance.save()
            return self.instance
        status_changed     = new_status     != self.original_status
        technician_changed = new_technician != self.original_technician
        if technician_changed:
            self.instance.technician = new_technician
        if status_changed:
            if commit:
                if technician_changed:
                    self.instance.save(update_fields=['technician'])
                try:
                    self.instance.change_status(new_status, changed_by=self.user)
                except ValidationError as e:
                    raise ValidationError({'status': str(e)})
        else:
            if commit:
                self.instance.save()
        return self.instance


class CommentForm(forms.ModelForm):
    class Meta:
        model = Comment
        fields = ['text']
        widgets = {
            'text': forms.Textarea(attrs={
                'rows': 3,
                'placeholder': 'Добавьте комментарий к заявке...',
                'class': 'form-control',
            }),
        }
        labels = {'text': ''}


class EquipmentForm(forms.ModelForm):
    class Meta:
        model = Equipment
        fields = ['model', 'serial', 'equipment_type', 'location',
                  'status', 'purchased_at', 'warranty_until', 'notes']
        widgets = {
            'purchased_at':   forms.DateInput(attrs={'type': 'date'}),
            'warranty_until': forms.DateInput(attrs={'type': 'date'}),
            'notes':          forms.Textarea(attrs={'rows': 3}),
        }


class UserEditForm(forms.Form):
    ROLE_CHOICES = [
        ('reporter',   'Пользователь'),
        ('dispatcher', 'Диспетчер'),
        ('technician', 'Техник'),
        ('admin',      'Администратор'),
    ]
    first_name = forms.CharField(label='Имя',      max_length=150, required=False)
    last_name  = forms.CharField(label='Фамилия',  max_length=150, required=False)
    email      = forms.EmailField(label='Email',   required=False)
    role       = forms.ChoiceField(label='Роль',   choices=ROLE_CHOICES)
    is_active  = forms.BooleanField(label='Активен', required=False)


class UserRegistrationForm(forms.Form):
    username  = forms.CharField(label='Имя пользователя', max_length=150,
        widget=forms.TextInput(attrs={
            'class': 'form-control', 'placeholder': 'Придумайте логин',
            'autocomplete': 'username'}))
    first_name = forms.CharField(label='Имя', max_length=150, required=False,
        widget=forms.TextInput(attrs={
            'class': 'form-control', 'placeholder': 'Иван'}))
    last_name  = forms.CharField(label='Фамилия', max_length=150, required=False,
        widget=forms.TextInput(attrs={
            'class': 'form-control', 'placeholder': 'Иванов'}))
    email      = forms.EmailField(label='Email', required=False,
        widget=forms.EmailInput(attrs={
            'class': 'form-control', 'placeholder': 'ivanov@mirmex.ru'}))
    password1  = forms.CharField(label='Пароль',
        widget=forms.PasswordInput(attrs={
            'class': 'form-control', 'placeholder': 'Минимум 8 символов',
            'autocomplete': 'new-password'}))
    password2  = forms.CharField(label='Подтверждение пароля',
        widget=forms.PasswordInput(attrs={
            'class': 'form-control', 'placeholder': 'Повторите пароль',
            'autocomplete': 'new-password'}))

    def clean_username(self):
        username = self.cleaned_data['username']
        if User.objects.filter(username=username).exists():
            raise ValidationError('Пользователь с таким именем уже существует.')
        return username

    def clean_password1(self):
        password = self.cleaned_data.get('password1')
        if password:
            validate_password(password)
        return password

    def clean(self):
        cleaned_data = super().clean()
        p1 = cleaned_data.get('password1')
        p2 = cleaned_data.get('password2')
        if p1 and p2 and p1 != p2:
            raise ValidationError({'password2': 'Пароли не совпадают.'})
        return cleaned_data


class ProfileForm(forms.Form):
    first_name = forms.CharField(label='Имя', max_length=150, required=False,
        widget=forms.TextInput(attrs={'class': 'form-control'}))
    last_name  = forms.CharField(label='Фамилия', max_length=150, required=False,
        widget=forms.TextInput(attrs={'class': 'form-control'}))
    email      = forms.EmailField(label='Email', required=False,
        widget=forms.EmailInput(attrs={'class': 'form-control'}))
'''
    add_code_block(doc, forms_full)

    # Сохраняем
    out_path = r'E:\vs_code(saves)\mirmex-helpdesk\.claude\worktrees\zen-elgamal\diploma_listing.docx'
    doc.save(out_path)
    print(f'Документ сохранён: {out_path}')


if __name__ == '__main__':
    build_doc()
